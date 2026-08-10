"""Step 2: fill Word document templates with project data.

Templates live under resources/templates/ and contain placeholder tokens
like {{job_number}}. fill_docx_template() copies the template to
output_path and replaces only those tokens - the rest of the document
(formatting, images, tables, headers/footers) is left untouched, so the
full template content never needs to be reproduced in code.
"""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.utils.logger import get_logger

logger = get_logger(__name__)

# Placeholder syntax used inside template documents, e.g. {{job_number}}.
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")

# Tick-box glyphs used for checkbox-style placeholders (see app/gui/constants.py).
UNCHECKED, CHECKED = "☐", "☒"
CHECKBOX_CHARS = (UNCHECKED, CHECKED)

# A run's font actually has four slots (ascii/hAnsi/eastAsia/cs - complex
# script); Word picks which slot governs a given character by its Unicode
# range, not always the plain "ascii" one python-docx's Font.name sets.
# The ballot-box glyphs above can end up routed through the eastAsia/cs
# slot, which may point at a theme font with no glyph for them (rendering
# as tofu/garbled). Segoe UI Symbol ships with Windows and covers them, so
# runs containing a checkbox glyph get all four slots forced to it.
SYMBOL_FONT = "Segoe UI Symbol"


def _force_symbol_font(run: Run, font_name: str = SYMBOL_FONT) -> None:
    run.font.name = font_name  # sets w:rFonts/@ascii and @hAnsi
    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    if rFonts.get(qn("w:hint")) is not None:
        del rFonts.attrib[qn("w:hint")]


def _run_text_offsets(runs: list[Run]) -> tuple[str, list[int]]:
    text = ""
    offsets = [0]
    for run in runs:
        text += run.text
        offsets.append(len(text))
    return text, offsets


def _run_range_for_span(offsets: list[int], start_char: int, end_char: int) -> tuple[int, int] | None:
    """Map a [start_char, end_char) text span back to the (start, end) run
    indices covering it exactly, or None if it doesn't align with run
    boundaries (i.e. a run only partially overlaps the span).
    """
    start_idx = end_idx = None
    for i, offset in enumerate(offsets[:-1]):
        if offset == start_char:
            start_idx = i
        if offsets[i + 1] == end_char:
            end_idx = i + 1
    if start_idx is None or end_idx is None:
        return None
    return start_idx, end_idx


def _replace_in_paragraph(
    paragraph: Paragraph, replacements: dict[str, str], skip_warning_keys: frozenset[str] = frozenset()
) -> None:
    """Replace each {{token}} with its value, touching only that token's
    own run(s). A paragraph can mix differently-formatted runs (e.g. one
    token highlighted, its neighbours not) - editing only the run(s) that
    actually spell out each token keeps every other run's formatting
    (highlight, bold, font...) exactly as the template author set it,
    instead of collapsing the whole paragraph into one run's formatting.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text, offsets = _run_text_offsets(runs)
    if "{{" not in full_text:
        return

    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    # Right-to-left: run indices computed from the initial (pre-edit) scan
    # stay valid for not-yet-processed matches, since edits only ever
    # remove runs at or after the current match.
    for match in reversed(matches):
        key = match.group(1)
        if key not in replacements:
            if key not in skip_warning_keys:
                logger.warning("No value provided for placeholder {{%s}}", key)
            continue

        rng = _run_range_for_span(offsets, match.start(), match.end())
        if rng is None:
            logger.warning("{{%s}} does not align with run boundaries, leaving as-is", key)
            continue
        start_idx, end_idx = rng

        target_runs = runs[start_idx:end_idx]
        value = replacements[key]
        target_runs[0].text = value
        if any(char in value for char in CHECKBOX_CHARS):
            _force_symbol_font(target_runs[0])
        for run in target_runs[1:]:
            run._element.getparent().remove(run._element)


def _set_checkbox_state(sdt, checked: bool) -> None:
    """Toggle a real Word checkbox content control (w:sdt/w14:checkbox):
    flips the stored checked flag and swaps the visible glyph to match, so
    it displays correctly even before Word recalculates the field.
    """
    sdtPr = sdt.find(qn("w:sdtPr"))
    checkbox_el = sdtPr.find(qn("w14:checkbox")) if sdtPr is not None else None
    if checkbox_el is None:
        return

    checked_el = checkbox_el.find(qn("w14:checked"))
    checked_el.set(qn("w14:val"), "1" if checked else "0")

    state_tag = qn("w14:checkedState") if checked else qn("w14:uncheckedState")
    state_el = checkbox_el.find(state_tag)
    glyph = chr(int(state_el.get(qn("w14:val")), 16))

    sdt_content = sdt.find(qn("w:sdtContent"))
    t_el = sdt_content.find(".//" + qn("w:t")) if sdt_content is not None else None
    if t_el is not None:
        t_el.text = glyph


def _iter_containers(document: Document):
    """The body plus every header/footer part - everywhere placeholders,
    checkboxes, or bullet lists might live."""
    yield document.element.body
    for section in document.sections:
        yield section.header._element
        yield section.footer._element


def _iter_all_paragraphs(document: Document):
    for container in _iter_containers(document):
        yield from (Paragraph(p, document) for p in container.iter(qn("w:p")))


def _iter_checkbox_controls(document: Document):
    for container in _iter_containers(document):
        yield from container.iter(qn("w:sdt"))


def _apply_checkbox_controls(document: Document, replacements: dict[str, str]) -> None:
    """Find real Word checkbox content controls tagged (via <w:tag>) with a
    replacements key, and toggle them to checked/unchecked based on whether
    that key's value is the CHECKED glyph - as opposed to the plain-text
    checkbox handling in _replace_in_paragraph, these are actual
    interactive Word checkboxes, so the tag is used instead of {{}} text.
    """
    for sdt in _iter_checkbox_controls(document):
        sdtPr = sdt.find(qn("w:sdtPr"))
        tag_el = sdtPr.find(qn("w:tag")) if sdtPr is not None else None
        if tag_el is None:
            continue
        tag = tag_el.get(qn("w:val"))
        if tag not in replacements:
            continue
        _set_checkbox_state(sdt, replacements[tag] == CHECKED)


# A "-" bulleted-list level definition matching the client's own template
# (a real Word numbering, not a typed dash character) - see
# _ensure_bullet_numbering().
BULLET_ABSTRACT_NUM_XML = """<w:abstractNum {nsdecls} w:abstractNumId="{abstract_id}">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="-"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Calibri" w:eastAsiaTheme="minorEastAsia" w:hAnsi="Calibri" w:cs="Calibri" w:hint="default"/></w:rPr>
    </w:lvl>
</w:abstractNum>"""


def _find_bullet_num_id(numbering_element) -> int | None:
    for num_el in numbering_element.findall(qn("w:num")):
        abstract_ref = num_el.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            continue
        abstract_el = numbering_element.find(
            f'{qn("w:abstractNum")}[@{qn("w:abstractNumId")}="{abstract_ref.get(qn("w:val"))}"]'
        )
        lvl0 = abstract_el.find(qn("w:lvl")) if abstract_el is not None else None
        lvl_text = lvl0.find(qn("w:lvlText")) if lvl0 is not None else None
        if lvl_text is not None and lvl_text.get(qn("w:val")) == "-":
            return int(num_el.get(qn("w:numId")))
    return None


def _ensure_bullet_numbering(document: Document) -> int:
    """Return the numId of a "-" bullet list definition in this document,
    reusing one if the template already has it (as the PS1 template does,
    inherited from the client's original), otherwise adding one.
    """
    numbering_element = document.part.numbering_part.element

    existing = _find_bullet_num_id(numbering_element)
    if existing is not None:
        return existing

    existing_abstract_ids = [
        int(el.get(qn("w:abstractNumId"))) for el in numbering_element.findall(qn("w:abstractNum"))
    ]
    existing_num_ids = [int(el.get(qn("w:numId"))) for el in numbering_element.findall(qn("w:num"))]
    new_abstract_id = max(existing_abstract_ids, default=-1) + 1
    new_num_id = max(existing_num_ids, default=0) + 1

    abstract_el = parse_xml(BULLET_ABSTRACT_NUM_XML.format(nsdecls=nsdecls("w"), abstract_id=new_abstract_id))
    num_el = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="{new_num_id}"><w:abstractNumId w:val="{new_abstract_id}"/></w:num>'
    )
    numbering_element.append(abstract_el)
    numbering_element.append(num_el)
    return new_num_id


def _insert_bullet_paragraph_after(
    anchor_element, parent, text: str, num_id: int, rPr_template=None, shd_template=None
):
    new_p = OxmlElement("w:p")
    anchor_element.addnext(new_p)
    new_paragraph = Paragraph(new_p, parent)
    run = new_paragraph.add_run(text)
    if rPr_template is not None:
        # rPr must be the run's first child per the schema.
        run._element.insert(0, deepcopy(rPr_template))
    numPr = new_paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = num_id
    if shd_template is not None:
        # w:shd follows w:numPr in the CT_PPrBase schema order - each
        # bullet paragraph gets its own shading, so a shaded block covers
        # exactly as many lines as there are bullets, growing/shrinking
        # with the list instead of a fixed-size decorative shape would.
        numPr.addnext(deepcopy(shd_template))
    return new_p


def _replace_placeholder_with_bullet_list(document: Document, field_name: str, lines: list[str]) -> bool:
    """Replace a {{field_name}} placeholder with one real bulleted-list
    paragraph per entry in `lines`, using the client template's own "-"
    bullet numbering rather than typed dash characters. Each new
    paragraph's run reuses the placeholder's own formatting (font, size,
    bold...) so it matches what the template author set up for it.
    """
    token = "{{" + field_name + "}}"

    for paragraph in _iter_all_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text, offsets = _run_text_offsets(runs)
        start_char = full_text.find(token)
        if start_char == -1:
            continue
        rng = _run_range_for_span(offsets, start_char, start_char + len(token))
        if rng is None:
            raise RuntimeError(f"{token!r} does not align with run boundaries in {full_text!r}")
        start_idx, end_idx = rng

        num_id = _ensure_bullet_numbering(document)
        template_rPr_el = runs[start_idx]._element.find(qn("w:rPr"))
        template_rPr = deepcopy(template_rPr_el) if template_rPr_el is not None else None
        pPr_el = paragraph._p.find(qn("w:pPr"))
        template_shd_el = pPr_el.find(qn("w:shd")) if pPr_el is not None else None
        template_shd = deepcopy(template_shd_el) if template_shd_el is not None else None

        for run in runs[start_idx:end_idx]:
            run._element.getparent().remove(run._element)

        anchor = paragraph._p
        for line in lines:
            anchor = _insert_bullet_paragraph_after(anchor, paragraph._parent, line, num_id, template_rPr, template_shd)

        if not paragraph.runs:
            paragraph._p.getparent().remove(paragraph._p)

        return True

    return False


# Paragraph style used for the top-level section headings that
# remove_unselected_sections() cuts the document on (see Specifications.docx
# and SPECIFICATION_SECTIONS in app/gui/constants.py).
SECTION_HEADING_STYLE = "Title 1"


def _remove_unselected_sections(document: Document, keep_titles: set[str]) -> None:
    """Delete every top-level section headed by a SECTION_HEADING_STYLE
    paragraph whose heading text isn't in `keep_titles` - the heading
    paragraph itself plus every paragraph/table up to (not including) the
    next such heading. Anything before the first heading (cover page,
    front matter) and the document's trailing `<w:sectPr>` are always kept.

    Word recalculates that heading style's own auto-numbering from
    whichever headings survive, so remaining sections don't need manual
    renumbering here.
    """
    body = document.element.body
    children = list(body)

    sect_pr_index = next((i for i, child in enumerate(children) if child.tag == qn("w:sectPr")), len(children))

    headings: list[tuple[int, str]] = []
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        paragraph = Paragraph(child, document)
        if paragraph.style is not None and paragraph.style.name == SECTION_HEADING_STYLE:
            headings.append((idx, paragraph.text.strip()))

    for i, (start_idx, title) in enumerate(headings):
        if title in keep_titles:
            continue
        end_idx = headings[i + 1][0] if i + 1 < len(headings) else sect_pr_index
        for child in children[start_idx:end_idx]:
            child.getparent().remove(child)


def _remove_unselected_table_rows(
    document: Document, table_index: int, keep_row_labels: set[str], header_rows: int = 1
) -> None:
    """Delete every data row (after the first `header_rows` rows) of
    `document.tables[table_index]` whose first-cell text isn't in
    `keep_row_labels` - e.g. the Material/Means of Compliance/Notes table
    in B2 Letter.docx, where the user picks which material rows apply.
    Row content itself is never touched, only whether the row survives.
    """
    table = document.tables[table_index]
    trs = table._tbl.findall(qn("w:tr"))
    for tr in trs[header_rows:]:
        tc = tr.find(qn("w:tc"))
        label = _Cell(tc, table).text.strip()
        if label not in keep_row_labels:
            tr.getparent().remove(tr)


def fill_docx_template(
    template_path: str,
    output_path: str,
    replacements: dict[str, Any],
    bullet_lists: dict[str, list[str]] | None = None,
    keep_sections: set[str] | None = None,
    keep_table_rows: dict[int, set[str]] | None = None,
) -> Path:
    """Copy template_path to output_path and replace {{placeholder}} tokens.

    replacements keys must match the placeholder names used in the
    template (without the surrounding {{ }}), e.g. {"job_number": "1234"}.
    bullet_lists keys work the same way, but each maps to a list of lines
    that become a genuine Word bulleted list (one paragraph per line) in
    place of the placeholder, instead of a plain-text substitution.
    keep_sections, if given, deletes every SECTION_HEADING_STYLE-headed
    section whose heading text isn't in the set (see
    _remove_unselected_sections) before any token replacement runs, so
    tokens that only exist inside a removed section never need a value.
    keep_table_rows, if given, maps a table index to the set of first-cell
    labels whose rows should survive in that table (see
    _remove_unselected_table_rows) - applied the same "before replacement"
    way as keep_sections.
    Raises FileNotFoundError if the template is missing, and
    FileExistsError if output_path already exists (never overwrites).
    """
    template = Path(template_path)
    output = Path(output_path)
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    if output.exists():
        raise FileExistsError(f"Output file already exists, not overwriting: {output}")

    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template, output)

    str_replacements = {key: "" if value is None else str(value) for key, value in replacements.items()}

    document = Document(output)

    if keep_sections is not None:
        _remove_unselected_sections(document, keep_sections)

    for table_index, keep_row_labels in (keep_table_rows or {}).items():
        _remove_unselected_table_rows(document, table_index, keep_row_labels)

    skip_warning_keys = frozenset((bullet_lists or {}).keys())

    for paragraph in _iter_all_paragraphs(document):
        _replace_in_paragraph(paragraph, str_replacements, skip_warning_keys)

    _apply_checkbox_controls(document, str_replacements)

    for field_name, lines in (bullet_lists or {}).items():
        _replace_placeholder_with_bullet_list(document, field_name, lines)

    document.save(output)
    logger.info("Filled template %s -> %s", template, output)
    return output
